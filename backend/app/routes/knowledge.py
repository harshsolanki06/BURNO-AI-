"""
BURNO AI OS — Knowledge Base Routes
POST   /api/knowledge/upload         — Upload file (PDF, TXT, MD, DOCX)
GET    /api/knowledge/documents      — List all documents
GET    /api/knowledge/search?q=...   — Full-text search across documents
GET    /api/knowledge/{id}           — Get single document detail
DELETE /api/knowledge/{id}           — Delete document
GET    /api/knowledge/stats          — Total docs, chunks, file types
"""
import io
import json
import re
from datetime import datetime, timezone
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc, func

from app.core.deps import get_db
from app.models.knowledge import KnowledgeDocument

router = APIRouter(prefix="/api/knowledge", tags=["Knowledge"])

# ─── Allowed file types ───────────────────────────────────────────────────────
ALLOWED_TYPES = {
    "application/pdf": "pdf",
    "text/plain": "txt",
    "text/markdown": "md",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
}
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".mdx", ".docx", ".doc"}
MAX_FILE_SIZE = 20 * 1024 * 1024   # 20 MB


# ─── Text extraction ─────────────────────────────────────────────────────────
def extract_text(filename: str, content: bytes) -> str:
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # PDF
    if ext == ".pdf":
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages:
                text = page.extract_text() or ""
                pages.append(text.strip())
            return "\n\n".join(p for p in pages if p)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"PDF extraction failed: {e}")

    # DOCX
    if ext in (".docx", ".doc"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paras)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"DOCX extraction failed: {e}")

    # Plain text / Markdown
    try:
        return content.decode("utf-8", errors="replace")
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Text decode failed: {e}")


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks for search and retrieval."""
    # Clean whitespace
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    if not text:
        return []

    # Try paragraph splitting first
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 2 <= chunk_size:
            current = (current + "\n\n" + para).strip() if current else para
        else:
            if current:
                chunks.append(current)
            # If single paragraph is larger than chunk_size, split by sentence
            if len(para) > chunk_size:
                words = para.split()
                sub = ""
                for word in words:
                    if len(sub) + len(word) + 1 <= chunk_size:
                        sub = (sub + " " + word).strip()
                    else:
                        if sub:
                            chunks.append(sub)
                        sub = word
                if sub:
                    current = sub
                else:
                    current = ""
            else:
                current = para

    if current:
        chunks.append(current)

    return [c for c in chunks if len(c.strip()) > 20]


def _serialize(doc: KnowledgeDocument, include_content: bool = False) -> dict:
    result = {
        "id": doc.id,
        "filename": doc.filename,
        "file_type": doc.file_type,
        "size_bytes": doc.size_bytes,
        "chunk_count": doc.chunk_count,
        "status": doc.status,
        "created_at": doc.created_at.isoformat() if doc.created_at else "",
        "preview": (doc.content or "")[:300] + ("…" if len(doc.content or "") > 300 else ""),
    }
    if include_content:
        result["content"] = doc.content
        result["chunks"] = doc.get_chunks()
    return result


# ─── Upload ──────────────────────────────────────────────────────────────────
@router.post("/upload")
async def upload_document(
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    # Validate extension
    filename = file.filename or "unknown"
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=f"File type '{ext}' not supported. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Read file
    raw = await file.read()
    if len(raw) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail=f"File too large. Max 20 MB.")
    if len(raw) == 0:
        raise HTTPException(status_code=400, detail="File is empty.")

    # Determine file type label
    file_type = ext.lstrip(".")

    # Extract text
    text = extract_text(filename, raw)
    if not text.strip():
        raise HTTPException(status_code=422, detail="Could not extract any text from this file.")

    # Chunk
    chunks = chunk_text(text)

    # Store in DB
    doc = KnowledgeDocument(
        filename=filename,
        file_type=file_type,
        content=text,
        chunks_json=json.dumps(chunks),
        size_bytes=len(raw),
        chunk_count=len(chunks),
        status="ready",
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    return {
        "id": doc.id,
        "filename": doc.filename,
        "file_type": doc.file_type,
        "size_bytes": doc.size_bytes,
        "chunk_count": len(chunks),
        "status": "ready",
        "preview": text[:300],
        "uploaded": True,
    }


# ─── List ─────────────────────────────────────────────────────────────────────
@router.get("/documents")
async def list_documents(
    limit: int = 50,
    offset: int = 0,
    file_type: Optional[str] = None,
    db: AsyncSession = Depends(get_db),
):
    stmt = select(KnowledgeDocument)
    if file_type and file_type != "all":
        stmt = stmt.where(KnowledgeDocument.file_type == file_type)
    stmt = stmt.order_by(desc(KnowledgeDocument.created_at)).offset(offset).limit(limit)
    result = await db.execute(stmt)
    docs = list(result.scalars().all())

    count_stmt = select(func.count(KnowledgeDocument.id))
    total = (await db.execute(count_stmt)).scalar() or 0

    type_stmt = select(KnowledgeDocument.file_type, func.count(KnowledgeDocument.id)).group_by(KnowledgeDocument.file_type)
    type_result = await db.execute(type_stmt)
    file_types = {row[0]: row[1] for row in type_result.all()}

    return {
        "documents": [_serialize(d) for d in docs],
        "total": total,
        "file_types": file_types,
    }


# ─── Search ──────────────────────────────────────────────────────────────────
@router.get("/search")
async def search_knowledge(
    q: str,
    limit: int = 20,
    db: AsyncSession = Depends(get_db),
):
    if not q.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    # Search in content (case-insensitive LIKE)
    stmt = (
        select(KnowledgeDocument)
        .where(KnowledgeDocument.content.ilike(f"%{q}%"))
        .order_by(desc(KnowledgeDocument.created_at))
        .limit(limit)
    )
    result = await db.execute(stmt)
    docs = list(result.scalars().all())

    # Extract matching chunks for each document
    hits = []
    query_lower = q.lower()
    for doc in docs:
        matching_chunks = []
        for chunk in doc.get_chunks():
            if query_lower in chunk.lower():
                # Highlight snippet: find position and return ±150 chars
                pos = chunk.lower().find(query_lower)
                start = max(0, pos - 100)
                end = min(len(chunk), pos + len(q) + 100)
                snippet = ("…" if start > 0 else "") + chunk[start:end] + ("…" if end < len(chunk) else "")
                matching_chunks.append(snippet)
                if len(matching_chunks) >= 3:
                    break
        hits.append({
            **_serialize(doc),
            "matching_chunks": matching_chunks,
            "match_count": len(matching_chunks),
        })

    return {
        "query": q,
        "results": hits,
        "count": len(hits),
    }


# ─── Get single ──────────────────────────────────────────────────────────────
@router.get("/{doc_id}")
async def get_document(doc_id: str, db: AsyncSession = Depends(get_db)):
    stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return _serialize(doc, include_content=True)


# ─── Delete ──────────────────────────────────────────────────────────────────
@router.delete("/{doc_id}")
async def delete_document(doc_id: str, db: AsyncSession = Depends(get_db)):
    stmt = select(KnowledgeDocument).where(KnowledgeDocument.id == doc_id)
    result = await db.execute(stmt)
    doc = result.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    await db.delete(doc)
    await db.commit()
    return {"deleted": True, "id": doc_id}


# ─── Stats ───────────────────────────────────────────────────────────────────
@router.get("/stats/summary")
async def knowledge_stats(db: AsyncSession = Depends(get_db)):
    total = (await db.execute(select(func.count(KnowledgeDocument.id)))).scalar() or 0
    chunks = (await db.execute(select(func.sum(KnowledgeDocument.chunk_count)))).scalar() or 0
    size = (await db.execute(select(func.sum(KnowledgeDocument.size_bytes)))).scalar() or 0
    type_result = await db.execute(
        select(KnowledgeDocument.file_type, func.count(KnowledgeDocument.id)).group_by(KnowledgeDocument.file_type)
    )
    return {
        "total_documents": total,
        "total_chunks": chunks,
        "total_size_bytes": size,
        "file_types": {row[0]: row[1] for row in type_result.all()},
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }
